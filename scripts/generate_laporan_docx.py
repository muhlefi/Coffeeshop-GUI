from datetime import date
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_PATH = "LAPORAN_TB_PBO_COFFEESHOP_2025_2026.docx"


def set_page_layout(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def apply_base_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color_hex, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color_hex)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.1


def add_title_page(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("LAPORAN TUGAS BESAR\nPEMROGRAMAN BERORIENTASI OBJEK")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(11, 37, 69)

    p2 = doc.add_paragraph("Sistem Manajemen Coffee Shop", style="Heading 2")
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run("\nSemester Genap 2025/2026\n").bold = True
    info.add_run(f"Tanggal Penyusunan Dokumen: {date.today().isoformat()}\n")
    info.add_run("Mata Kuliah: Pemrograman Berorientasi Objek")

    doc.add_paragraph()
    table = doc.add_table(rows=7, cols=2)
    table.style = "Table Grid"
    rows = [
        ("Kelas", "................"),
        ("Nomor Kelompok", "................"),
        ("Topik", "Coffee Shop"),
        ("Nama Project", "X_Y_CoffeeShop"),
        ("Database", "X_Y_CoffeeShop"),
        ("Dosen Pengampu", "................"),
        ("Asisten Praktikum", "................"),
    ]
    for i, (k, v) in enumerate(rows):
        table.cell(i, 0).text = k
        table.cell(i, 1).text = v

    doc.add_page_break()


def add_member_table(doc: Document) -> None:
    doc.add_heading("Identitas Kelompok", level=1)
    t = doc.add_table(rows=5, cols=3)
    t.style = "Table Grid"
    t.cell(0, 0).text = "No"
    t.cell(0, 1).text = "Nama Lengkap"
    t.cell(0, 2).text = "NPM"
    for i in range(1, 5):
        t.cell(i, 0).text = str(i)
        t.cell(i, 1).text = "................"
        t.cell(i, 2).text = "................"


def add_intro(doc: Document) -> None:
    doc.add_heading("Bab 1 Pendahuluan", level=1)
    doc.add_heading("1.1 Latar Belakang", level=2)
    doc.add_paragraph(
        "Coffee shop modern membutuhkan sistem yang mampu mengelola data master, transaksi, "
        "dan pelaporan secara cepat serta konsisten. Proses manual berisiko menimbulkan kesalahan "
        "input, keterlambatan rekap, dan inkonsistensi stok."
    )
    doc.add_paragraph(
        "Aplikasi GUI Java ini dirancang untuk membantu pengelolaan operasional harian melalui "
        "fitur master data, transaksi penjualan, transaksi pembelian/restock, dan laporan periodik."
    )

    doc.add_heading("1.2 Rumusan Masalah", level=2)
    for text in [
        "Bagaimana membangun aplikasi desktop Java GUI untuk operasional coffee shop berbasis MySQL?",
        "Bagaimana menerapkan prinsip PBO secara utuh pada implementasi kode?",
        "Bagaimana menyediakan data report yang akurat untuk monitoring bisnis?"
    ]:
        doc.add_paragraph(text, style="List Bullet")

    doc.add_heading("1.3 Tujuan", level=2)
    for text in [
        "Membangun aplikasi GUI Java yang user-friendly dan stabil.",
        "Menerapkan konsep PBO: encapsulation, inheritance, polymorphism, abstraction, interface/abstract class.",
        "Menyediakan fitur CRUD master data, transaksi, dan report sesuai ketentuan tugas."
    ]:
        doc.add_paragraph(text, style="List Number")

    doc.add_heading("1.4 Batasan Masalah", level=2)
    for text in [
        "Aplikasi berbasis desktop Java Swing dengan database MySQL.",
        "Fokus pada modul master data, transaksi penjualan, transaksi pembelian, dan report.",
        "Belum mencakup integrasi pembayaran online dan multi-cabang."
    ]:
        doc.add_paragraph(text, style="List Bullet")


def add_analysis_design(doc: Document) -> None:
    doc.add_heading("Bab 2 Analisis dan Perancangan", level=1)
    doc.add_heading("2.1 Kebutuhan Fungsional", level=2)
    for text in [
        "Pengelolaan master produk (create, read, update, delete, search).",
        "Pengelolaan master customer (create, read, update, delete, search).",
        "Transaksi penjualan (header-detail, hitung diskon/pajak/total, update stok keluar).",
        "Transaksi pembelian (header-detail, update stok masuk).",
        "Laporan penjualan bulanan, produk terlaris, dan pergerakan stok."
    ]:
        doc.add_paragraph(text, style="List Bullet")

    doc.add_heading("2.2 Use Case Ringkas", level=2)
    uc = doc.add_table(rows=6, cols=3)
    uc.style = "Table Grid"
    uc.cell(0, 0).text = "No"
    uc.cell(0, 1).text = "Aktor"
    uc.cell(0, 2).text = "Use Case"
    use_cases = [
        ("1", "Admin", "Kelola master produk dan customer"),
        ("2", "Kasir", "Mencatat transaksi penjualan"),
        ("3", "Gudang", "Mencatat transaksi pembelian/restock"),
        ("4", "Manajer", "Melihat laporan penjualan"),
        ("5", "Manajer", "Melihat laporan produk terlaris"),
    ]
    for i, row in enumerate(use_cases, start=1):
        uc.cell(i, 0).text = row[0]
        uc.cell(i, 1).text = row[1]
        uc.cell(i, 2).text = row[2]

    doc.add_heading("2.3 Entity Relationship Diagram (ERD)", level=2)
    doc.add_paragraph(
        "Tempelkan gambar ERD final pada bagian ini. ERD wajib menunjukkan relasi tabel: "
        "users, customers, suppliers, categories, products, sales, sales_detail, purchases, purchase_detail."
    )
    doc.add_paragraph("[TEMPAT GAMBAR ERD]", style="Intense Quote")

    doc.add_heading("2.4 Class Diagram (UML)", level=2)
    doc.add_paragraph(
        "Tempelkan gambar class diagram final pada bagian ini. Pastikan terlihat relasi inheritance, "
        "interface implementation, dan layer UI-Service-Repository."
    )
    doc.add_paragraph("[TEMPAT GAMBAR CLASS DIAGRAM UML]", style="Intense Quote")


def add_implementation(doc: Document) -> None:
    doc.add_heading("Bab 3 Implementasi", level=1)
    doc.add_heading("3.1 Arsitektur dan Struktur Package", level=2)
    for text in [
        "com.coffeeshop.config",
        "com.coffeeshop.model",
        "com.coffeeshop.repository dan com.coffeeshop.repository.impl",
        "com.coffeeshop.service",
        "com.coffeeshop.ui dan com.coffeeshop.ui.panel",
        "com.coffeeshop.report, com.coffeeshop.exception, com.coffeeshop.util",
    ]:
        doc.add_paragraph(text, style="List Bullet")

    doc.add_heading("3.2 Implementasi Fitur Utama", level=2)
    features = doc.add_table(rows=6, cols=3)
    features.style = "Table Grid"
    features.cell(0, 0).text = "No"
    features.cell(0, 1).text = "Modul"
    features.cell(0, 2).text = "Deskripsi"
    rows = [
        ("1", "Dashboard", "Test koneksi database dan ringkasan awal."),
        ("2", "Master Produk", "CRUD + Search produk."),
        ("3", "Master Customer", "CRUD + Search customer."),
        ("4", "Transaksi Penjualan", "Simpan sales + sales_detail + stok keluar."),
        ("5", "Transaksi Pembelian", "Simpan purchases + purchase_detail + stok masuk."),
    ]
    for i, row in enumerate(rows, start=1):
        features.cell(i, 0).text = row[0]
        features.cell(i, 1).text = row[1]
        features.cell(i, 2).text = row[2]

    doc.add_heading("3.3 Penerapan Prinsip PBO pada Kode", level=2)
    pbo = doc.add_table(rows=10, cols=3)
    pbo.style = "Table Grid"
    pbo.cell(0, 0).text = "Prinsip"
    pbo.cell(0, 1).text = "Implementasi"
    pbo.cell(0, 2).text = "Contoh Class/File"
    values = [
        ("Encapsulation", "Atribut private + getter/setter.", "model/Product.java, model/Customer.java"),
        ("Inheritance", "Class turunan dari BaseEntity.", "model/BaseEntity.java -> Product/Customer/Sale/Purchase"),
        ("Polymorphism", "Implementasi interface report generator.", "report/ReportGenerator.java, MonthlySalesReport.java"),
        ("Abstraction", "Abstract class untuk logika transaksi.", "service/AbstractTransactionService.java"),
        ("Class dan Object", "Semua entity direpresentasi class.", "model/*.java"),
        ("Constructor", "Constructor default/parameter pada model.", "model/Product.java, model/Customer.java"),
        ("Access Modifier", "public/private/protected pada layer kode.", "seluruh package"),
        ("Interface", "Kontrak repository dan report.", "repository/*.java, report/ReportGenerator.java"),
        ("User-defined Exception", "Validasi input kosong & data invalid.", "exception/InputKosongException.java"),
    ]
    for i, row in enumerate(values, start=1):
        pbo.cell(i, 0).text = row[0]
        pbo.cell(i, 1).text = row[1]
        pbo.cell(i, 2).text = row[2]

    doc.add_heading("3.4 Exception Handling", level=2)
    for text in [
        "InputKosongException: dipakai untuk validasi field wajib.",
        "DataTidakValidException: dipakai untuk validasi nilai data/ID/logic bisnis.",
        "NumberFormatException: menangani input angka yang tidak valid.",
        "SQLException: menangani error query/koneksi database."
    ]:
        doc.add_paragraph(text, style="List Bullet")


def add_testing(doc: Document) -> None:
    doc.add_heading("Bab 4 Pengujian", level=1)
    doc.add_paragraph("Pengujian dilakukan dengan metode uji fungsional berbasis skenario.")

    test = doc.add_table(rows=9, cols=5)
    test.style = "Table Grid"
    headers = ["No", "Skenario Uji", "Input", "Hasil Diharapkan", "Status"]
    for i, header in enumerate(headers):
        test.cell(0, i).text = header

    scenarios = [
        ("1", "Tambah produk baru", "Data produk valid", "Data tersimpan dan tampil di tabel"),
        ("2", "Update data customer", "Edit customer valid", "Data ter-update"),
        ("3", "Simpan transaksi penjualan", "Item + qty valid", "Sales tersimpan, stok berkurang"),
        ("4", "Simpan transaksi pembelian", "Item + qty valid", "Purchase tersimpan, stok bertambah"),
        ("5", "Laporan penjualan bulanan", "Pilih tahun", "Data agregat bulanan tampil"),
        ("6", "Laporan produk terlaris", "Pilih periode", "Top produk tampil sesuai qty"),
        ("7", "Input kosong wajib", "Kosongkan field wajib", "Muncul InputKosongException"),
        ("8", "Input non-angka", "Isi huruf pada field angka", "Muncul NumberFormatException"),
    ]
    for i, row in enumerate(scenarios, start=1):
        for j, val in enumerate(row):
            test.cell(i, j).text = val
        test.cell(i, 4).text = "Lulus / Belum"

    doc.add_paragraph("Catatan: isi status akhir berdasarkan hasil pengujian aktual saat demo final.")


def add_conclusion(doc: Document) -> None:
    doc.add_heading("Bab 5 Kesimpulan dan Saran", level=1)
    doc.add_heading("5.1 Kesimpulan", level=2)
    for text in [
        "Aplikasi berhasil memenuhi syarat tugas besar PBO dengan penerapan konsep OOP yang relevan.",
        "Fitur master data, transaksi, dan report telah terintegrasi dengan database MySQL.",
        "Exception handling dan validasi input meningkatkan stabilitas aplikasi."
    ]:
        doc.add_paragraph(text, style="List Bullet")

    doc.add_heading("5.2 Saran Pengembangan", level=2)
    for text in [
        "Menambahkan autentikasi login berbasis role yang lebih ketat.",
        "Menambahkan ekspor report ke PDF/Excel.",
        "Menambahkan dashboard statistik grafis untuk analisis penjualan."
    ]:
        doc.add_paragraph(text, style="List Bullet")


def add_appendix(doc: Document) -> None:
    doc.add_heading("Lampiran", level=1)
    doc.add_heading("A. Checklist Screenshot untuk Laporan", level=2)
    for text in [
        "Halaman Dashboard",
        "Form Master Produk (data terisi)",
        "Form Master Customer (data terisi)",
        "Transaksi Penjualan (item dan total terlihat)",
        "Transaksi Pembelian (item dan total terlihat)",
        "Report Penjualan Bulanan",
        "Report Produk Terlaris",
        "Report Pergerakan Stok",
        "Pesan error custom exception (input kosong)",
        "Pesan error default exception (input angka tidak valid)",
    ]:
        doc.add_paragraph(text, style="List Number")

    doc.add_heading("B. Daftar File Pengumpulan", level=2)
    checklist = doc.add_table(rows=6, cols=3)
    checklist.style = "Table Grid"
    checklist.cell(0, 0).text = "No"
    checklist.cell(0, 1).text = "Nama File"
    checklist.cell(0, 2).text = "Keterangan"
    rows = [
        ("1", "X_Y_Topik (Project NetBeans/Java)", "Project source code"),
        ("2", "X_Y_Topik.sql", "Export database"),
        ("3", "X_Y_Topik.docx", "Dokumen laporan"),
        ("4", "X_Y_Topik.txt", "Identitas anggota"),
        ("5", "X_Y_Topik_TB.zip", "Paket final pengumpulan"),
    ]
    for i, row in enumerate(rows, start=1):
        checklist.cell(i, 0).text = row[0]
        checklist.cell(i, 1).text = row[1]
        checklist.cell(i, 2).text = row[2]


def set_table_cell_margins(table):
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tc_pr = tc.get_or_add_tcPr()
            tc_mar = OxmlElement("w:tcMar")
            for side in ("top", "start", "bottom", "end"):
                node = OxmlElement(f"w:{side}")
                node.set(qn("w:w"), "120")
                node.set(qn("w:type"), "dxa")
                tc_mar.append(node)
            tc_pr.append(tc_mar)


def enforce_table_format(doc: Document) -> None:
    for table in doc.tables:
        set_table_cell_margins(table)


def main() -> None:
    doc = Document()
    set_page_layout(doc)
    apply_base_styles(doc)
    add_title_page(doc)
    add_member_table(doc)
    add_intro(doc)
    add_analysis_design(doc)
    add_implementation(doc)
    add_testing(doc)
    add_conclusion(doc)
    add_appendix(doc)
    enforce_table_format(doc)
    doc.save(OUTPUT_PATH)
    print(f"Generated: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
